"""PDF & DOCX Document Parser Module.

PyPDF2 ile PDF parsing ve python-docx ile DOCX parsing.
- Metin cikarimi
- Metadata cikarimi
- Sayfa/bolum bazli chunk'lar
- Semantic chunking entegrasyonu
"""

import hashlib
import io
import tempfile
import time
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

import structlog

logger = structlog.get_logger(__name__)


@dataclass
class ParsedDocument:
    """Parse edilmis dokuman veri yapisi."""

    source_type: str  # "pdf" veya "docx"
    filename: str = ""
    title: str = ""
    author: str = ""
    subject: str = ""
    keywords: str = ""
    created_date: str = ""
    modified_date: str = ""
    page_count: int = 0
    word_count: int = 0
    content: str = ""
    pages: List[Dict[str, Any]] = field(default_factory=list)
    sections: List[Dict[str, Any]] = field(default_factory=list)
    tables: List[List[str]] = field(default_factory=list)
    content_hash: str = ""
    parse_time_ms: float = 0.0

    def to_knowledge_dict(self, company_id: int, branch_id: Optional[int] = None) -> Dict[str, Any]:
        """KnowledgeBase kayit formatina donustur."""
        return {
            "company_id": company_id,
            "branch_id": branch_id,
            "source_type": self.source_type,
            "source_title": self.title or self.filename,
            "source_description": self.subject or "",
            "raw_content": self.content,
            "raw_content_hash": self.content_hash,
            "content_metadata": {
                "filename": self.filename,
                "author": self.author,
                "keywords": self.keywords,
                "page_count": self.page_count,
                "word_count": self.word_count,
                "created_date": self.created_date,
                "modified_date": self.modified_date,
                "tables": self.tables,
            }
        }


class PDFParser:
    """PDF parser sinifi.

    PyPDF2 ile PDF dosyalarindan metin ve metadata cikarir.
    Sayfa bazli chunk'lar uretir.
    """

    def __init__(self, extract_tables: bool = True) -> None:
        self.extract_tables = extract_tables

    def parse(self, file_content: Union[bytes, str]) -> ParsedDocument:
        """PDF'i parse et.

        Args:
            file_content: PDF dosya icerigi (bytes veya dosya yolu).

        Returns:
            ParsedDocument veri yapisi.
        """
        start_time = time.time()

        try:
            from PyPDF2 import PdfReader
        except ImportError:
            logger.error("PyPDF2 not installed")
            return ParsedDocument(source_type="pdf", content="")

        doc = ParsedDocument(source_type="pdf")

        try:
            if isinstance(file_content, str):
                reader = PdfReader(file_content)
            else:
                reader = PdfReader(io.BytesIO(file_content))

            # Metadata cikar
            meta = reader.metadata
            if meta:
                doc.title = str(meta.get("/Title", "")) if meta.get("/Title") else ""
                doc.author = str(meta.get("/Author", "")) if meta.get("/Author") else ""
                doc.subject = str(meta.get("/Subject", "")) if meta.get("/Subject") else ""
                doc.keywords = str(meta.get("/Keywords", "")) if meta.get("/Keywords") else ""
                doc.created_date = str(meta.get("/CreationDate", "")) if meta.get("/CreationDate") else ""
                doc.modified_date = str(meta.get("/ModDate", "")) if meta.get("/ModDate") else ""

            # Sayfa sayisi
            doc.page_count = len(reader.pages)

            # Sayfa bazli metin cikarimi
            full_text_parts = []
            for i, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text() or ""
                except Exception as e:
                    logger.warning(f"page_extract_error", page=i, error=str(e))
                    page_text = ""

                if page_text.strip():
                    page_dict = {
                        "page_number": i + 1,
                        "content": page_text.strip(),
                        "char_count": len(page_text),
                    }
                    doc.pages.append(page_dict)
                    full_text_parts.append(page_text)

            # Tam metin
            doc.content = "\n\n".join(full_text_parts)
            doc.word_count = len(doc.content.split())
            doc.content_hash = hashlib.sha256(doc.content.encode("utf-8")).hexdigest()

        except Exception as e:
            logger.error("pdf_parse_error", error=str(e))
            doc.content = ""

        doc.parse_time_ms = (time.time() - start_time) * 1000
        logger.info(
            "pdf_parsed",
            pages=doc.page_count,
            words=doc.word_count,
            time_ms=round(doc.parse_time_ms, 2),
        )

        return doc

    def parse_file(self, file_path: str) -> ParsedDocument:
        """Dosya yolundan PDF parse et."""
        return self.parse(file_path)


class DOCXParser:
    """DOCX parser sinifi.

    python-docx ile DOCX dosyalarindan metin, baslik,
    tablo ve metadata cikarir. Bolum bazli chunk'lar uretir.
    """

    def __init__(self, extract_tables: bool = True, extract_styles: bool = True) -> None:
        self.extract_tables = extract_tables
        self.extract_styles = extract_styles

    def parse(self, file_content: Union[bytes, str]) -> ParsedDocument:
        """DOCX'i parse et.

        Args:
            file_content: DOCX dosya icerigi (bytes veya dosya yolu).

        Returns:
            ParsedDocument veri yapisi.
        """
        start_time = time.time()

        try:
            import docx
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
        except ImportError:
            logger.error("python-docx not installed")
            return ParsedDocument(source_type="docx", content="")

        doc = ParsedDocument(source_type="docx")

        try:
            if isinstance(file_content, str):
                document = docx.Document(file_content)
            else:
                document = docx.Document(io.BytesIO(file_content))

            # Core properties
            core_props = document.core_properties
            doc.title = core_props.title or ""
            doc.author = core_props.author or ""
            doc.subject = core_props.subject or ""
            doc.keywords = core_props.keywords or ""
            doc.created_date = str(core_props.created) if core_props.created else ""
            doc.modified_date = str(core_props.modified) if core_props.modified else ""

            # Paragraf bazli metin cikarimi
            full_text_parts = []
            current_section: Dict[str, Any] = {"heading": "", "content": "", "level": 0}
            sections: List[Dict[str, Any]] = []

            for para in document.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # Baslik tespiti
                if para.style.name.startswith("Heading"):
                    # Onceki bolumu kaydet
                    if current_section["content"]:
                        sections.append(current_section.copy())

                    level = 1
                    try:
                        level = int(para.style.name.replace("Heading ", ""))
                    except ValueError:
                        level = 1

                    current_section = {
                        "heading": text,
                        "content": "",
                        "level": level,
                    }
                    doc.pages.append({
                        "type": "heading",
                        "content": text,
                        "level": level,
                    })
                else:
                    current_section["content"] += text + "\n"
                    full_text_parts.append(text)

            # Son bolumu kaydet
            if current_section["content"]:
                sections.append(current_section)

            doc.sections = [
                {
                    "heading": s["heading"],
                    "content": s["content"].strip(),
                    "level": s["level"],
                    "word_count": len(s["content"].split()),
                }
                for s in sections
            ]

            # Tablolari cikar
            if self.extract_tables:
                for table in document.tables:
                    table_rows = []
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells]
                        if any(row_text):
                            table_rows.append(" | ".join(row_text))
                    if table_rows:
                        doc.tables.append(table_rows)
                        # Tablo metnini de ekle
                        full_text_parts.append("\n".join(table_rows))

            # Sayfa sayisi (yaklasik - paragraf bazli)
            doc.page_count = len(doc.sections) if doc.sections else max(1, len(full_text_parts) // 20)

            # Tam metin
            doc.content = "\n\n".join(full_text_parts)
            doc.word_count = len(doc.content.split())
            doc.content_hash = hashlib.sha256(doc.content.encode("utf-8")).hexdigest()

        except Exception as e:
            logger.error("docx_parse_error", error=str(e))
            doc.content = ""

        doc.parse_time_ms = (time.time() - start_time) * 1000
        logger.info(
            "docx_parsed",
            sections=len(doc.sections),
            words=doc.word_count,
            tables=len(doc.tables),
            time_ms=round(doc.parse_time_ms, 2),
        )

        return doc

    def parse_file(self, file_path: str) -> ParsedDocument:
        """Dosya yolundan DOCX parse et."""
        return self.parse(file_path)


class DocumentParser:
    """Unified document parser.

    Dosya tipine gore otomatik parser secimi.
    """

    SUPPORTED_TYPES = {
        ".pdf": PDFParser,
        ".docx": DOCXParser,
    }

    @classmethod
    def parse(
        cls,
        file_content: Union[bytes, str],
        filename: str = "",
        **kwargs: Any,
    ) -> ParsedDocument:
        """Dosyayi parse et (tip otomatik tespit).

        Args:
            file_content: Dosya icerigi.
            filename: Dosya adi (uzanti tespiti icin).

        Returns:
            ParsedDocument.
        """
        # Dosya tipi tespiti
        file_type = None
        if filename:
            ext = Path(filename).suffix.lower()
            file_type = ext if ext in cls.SUPPORTED_TYPES else None

        # Magic bytes ile tespit (filename yoksa)
        if not file_type and isinstance(file_content, bytes):
            if file_content[:4] == b"%PDF":
                file_type = ".pdf"
            elif file_content[:4] == b"PK\x03\x04":
                file_type = ".docx"

        if not file_type:
            logger.error("unsupported_file_type", filename=filename)
            return ParsedDocument(source_type="unknown", filename=filename, content="")

        # Parser sec ve calistir
        parser_class = cls.SUPPORTED_TYPES[file_type]
        parser = parser_class(**kwargs)
        doc = parser.parse(file_content)
        doc.filename = filename

        return doc

    @classmethod
    def is_supported(cls, filename: str) -> bool:
        """Dosya tipi destekleniyor mu?"""
        ext = Path(filename).suffix.lower()
        return ext in cls.SUPPORTED_TYPES


# =============================================================================
# Async Wrappers
# =============================================================================

async def parse_document_async(
    file_content: Union[bytes, str],
    filename: str,
    company_id: int,
    branch_id: Optional[int] = None,
    **kwargs: Any,
) -> Dict[str, Any]:
    """Async dokuman parse et ve knowledge dict dondur.

    Celery task icin kullanilir.

    Returns:
        KnowledgeBase kayit formatinda dict.
    """
    import asyncio

    loop = asyncio.get_event_loop()
    doc = await loop.run_in_executor(
        None,
        lambda: DocumentParser.parse(file_content, filename, **kwargs),
    )
    return doc.to_knowledge_dict(company_id, branch_id)


async def parse_upload_file(
    file: Any,  # UploadFile
    company_id: int,
    branch_id: Optional[int] = None,
) -> Dict[str, Any]:
    """UploadFile'dan dokuman parse et.

    FastAPI UploadFile ile kullanilir.

    Returns:
        KnowledgeBase kayit formatinda dict.
    """
    content = await file.read()
    filename = getattr(file, "filename", "unknown")
    return await parse_document_async(
        content, filename, company_id, branch_id
    )
