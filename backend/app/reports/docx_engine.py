"""
DOCX export engine using python-docx.

Generates professional Word documents with company-branded headers,
footers, logos, styled tables, and section formatting.
"""

import os
from datetime import datetime
from io import BytesIO
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.shared import Inches, Pt, RGBColor, Cm, Mm
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from app.reports.constants import DEFAULT_TEMPLATE_CONFIG


class DOCXReportEngine:
    """
    Professional DOCX report generator using python-docx.

    Supports branded headers/footers, tables, styled sections,
    and company logo integration.
    """

    def __init__(self, template_config: Optional[Dict[str, Any]] = None):
        self.config = {**DEFAULT_TEMPLATE_CONFIG, **(template_config or {})}
        self.colors = self._parse_colors()
        self.logo_path: Optional[str] = None

    def _parse_colors(self) -> Dict[str, RGBColor]:
        """Parse hex colors to RGBColor objects."""

        def hex_to_rgb(hex_str: str) -> RGBColor:
            hex_str = hex_str.lstrip("#")
            return RGBColor(
                int(hex_str[0:2], 16),
                int(hex_str[2:4], 16),
                int(hex_str[4:6], 16),
            )

        return {
            "primary": hex_to_rgb(self.config.get("primary_color", "#2563EB")),
            "secondary": hex_to_rgb(self.config.get("secondary_color", "#1E40AF")),
            "accent": hex_to_rgb(self.config.get("accent_color", "#3B82F6")),
            "text": hex_to_rgb(self.config.get("text_color", "#1F2937")),
            "light_bg": hex_to_rgb(self.config.get("light_bg", "#F3F4F6")),
            "white": hex_to_rgb(self.config.get("white", "#FFFFFF")),
        }

    def set_logo(self, logo_path: Optional[str]) -> None:
        """Set company logo path."""
        self.logo_path = logo_path

    def _hex_to_shd(self, hex_str: str) -> str:
        """Convert hex color to Word shading value."""
        return hex_str.lstrip("#")

    def generate(
        self,
        output_path: str,
        title: str,
        subtitle: Optional[str] = None,
        sections: Optional[List[Dict[str, Any]]] = None,
        tables: Optional[List[Dict[str, Any]]] = None,
        summary_stats: Optional[List[Dict[str, Any]]] = None,
    ) -> str:
        """
        Generate a DOCX report.

        Args:
            output_path: File path for output.
            title: Document title.
            subtitle: Optional subtitle.
            sections: Text sections [{"heading": str, "content": str|list}].
            tables: Data tables [{"title": str, "headers": [...], "rows": [...]}].
            summary_stats: Key-value stats [{"label": str, "value": str}].

        Returns:
            output_path on success.
        """
        doc = Document()

        # Page orientation
        section = doc.sections[0]
        if self.config.get("orientation") == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
            new_width = section.page_height
            section.page_width = new_width
            section.page_height = section.page_width

        # Margins
        margin_top = self.config.get("margin_top", 72)
        margin_bottom = self.config.get("margin_bottom", 72)
        margin_left = self.config.get("margin_left", 72)
        margin_right = self.config.get("margin_right", 72)
        section.top_margin = Pt(margin_top)
        section.bottom_margin = Pt(margin_bottom)
        section.left_margin = Pt(margin_left)
        section.right_margin = Pt(margin_right)

        # --- Header ---
        self._add_header(section, title)

        # --- Footer ---
        self._add_footer(section)

        # --- Title ---
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(self.config.get("title_font_size", 18))
        title_run.font.color.rgb = self.colors["primary"]
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.space_after = Pt(8)

        if subtitle:
            sub_para = doc.add_paragraph()
            sub_run = sub_para.add_run(subtitle)
            sub_run.font.size = Pt(self.config.get("subtitle_font_size", 12))
            sub_run.font.color.rgb = self.colors["text"]
            sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_para.space_after = Pt(16)

        # Generated date
        gen_para = doc.add_paragraph()
        gen_run = gen_para.add_run(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
        gen_run.font.size = Pt(8)
        gen_run.font.color.rgb = RGBColor(128, 128, 128)
        gen_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        gen_para.space_after = Pt(16)

        # --- Summary Stats ---
        if summary_stats:
            self._add_stats_section(doc, summary_stats)

        # --- Sections ---
        if sections:
            for sec in sections:
                self._add_text_section(doc, sec)

        # --- Tables ---
        if tables:
            for table_data in tables:
                self._add_data_table(doc, table_data)

        doc.save(output_path)
        return output_path

    def _add_header(self, section: Any, title: str) -> None:
        """Add branded header with logo."""
        header = section.header
        header.is_linked_to_previous = False

        # Clear existing
        for para in header.paragraphs:
            para.clear()

        if not header.paragraphs:
            header.add_paragraph()

        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Add logo if available
        if self.logo_path and os.path.exists(self.logo_path):
            try:
                run = header_para.add_run()
                run.add_picture(self.logo_path, width=Mm(30))
            except Exception:
                pass

        # Title in header
        run = header_para.add_run(f"  |  {title}")
        run.font.size = Pt(8)
        run.font.color.rgb = self.colors["primary"]
        run.bold = True

        # Add bottom border to header
        pPr = header_para._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="{self.config.get("primary_color", "#2563EB").lstrip("#")}"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

    def _add_footer(self, section: Any) -> None:
        """Add branded footer with page numbers."""
        footer = section.footer
        footer.is_linked_to_previous = False

        if not footer.paragraphs:
            footer.add_paragraph()

        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Footer text
        footer_text = self.config.get("footer_text", "Generated by Report Engine")
        run = footer_para.add_run(footer_text)
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor(128, 128, 128)

        # Add top border
        pPr = footer_para._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:space="1" w:color="E5E7EB"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

    def _add_stats_section(self, doc: Document, stats: List[Dict[str, Any]]) -> None:
        """Add a summary statistics row as a styled table."""
        table = doc.add_table(rows=2, cols=len(stats))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row (labels)
        for i, stat in enumerate(stats):
            cell = table.rows[0].cells[i]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(str(stat.get("label", "")))
            run.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = self.colors["white"]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Background color
            shading = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="{self.config.get("primary_color", "#2563EB").lstrip("#")}"/>'
            )
            cell._tc.get_or_add_tcPr().append(shading)

        # Value row
        for i, stat in enumerate(stats):
            cell = table.rows[1].cells[i]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(str(stat.get("value", "")))
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = self.colors["primary"]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            shading = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="{self.config.get("light_bg", "#F3F4F6").lstrip("#")}"/>'
            )
            cell._tc.get_or_add_tcPr().append(shading)

        doc.add_paragraph()  # Spacer

    def _add_text_section(self, doc: Document, section: Dict[str, Any]) -> None:
        """Add a text section with heading and content."""
        heading = section.get("heading")
        content = section.get("content", "")

        if heading:
            h_para = doc.add_paragraph()
            h_run = h_para.add_run(heading)
            h_run.bold = True
            h_run.font.size = Pt(14)
            h_run.font.color.rgb = self.colors["primary"]
            h_para.space_after = Pt(6)

        if isinstance(content, list):
            for para_text in content:
                p = doc.add_paragraph()
                run = p.add_run(str(para_text))
                run.font.size = Pt(self.config.get("body_font_size", 9))
                run.font.color.rgb = self.colors["text"]
                p.space_after = Pt(4)
        elif content:
            p = doc.add_paragraph()
            run = p.add_run(str(content))
            run.font.size = Pt(self.config.get("body_font_size", 9))
            run.font.color.rgb = self.colors["text"]
            p.space_after = Pt(8)

    def _add_data_table(self, doc: Document, table_data: Dict[str, Any]) -> None:
        """Add a styled data table."""
        title = table_data.get("title")
        headers = table_data.get("headers", [])
        rows = table_data.get("rows", [])

        if title:
            t_para = doc.add_paragraph()
            t_run = t_para.add_run(title)
            t_run.bold = True
            t_run.font.size = Pt(12)
            t_run.font.color.rgb = self.colors["text"]
            t_para.space_after = Pt(6)

        if not headers or not rows:
            doc.add_paragraph("No data available.").runs[0].font.color.rgb = RGBColor(128, 128, 128)
            return

        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        # Header row
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(str(header))
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = self.colors["white"]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            shading = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="{self.config.get("primary_color", "#2563EB").lstrip("#")}"/>'
            )
            cell._tc.get_or_add_tcPr().append(shading)

        # Data rows
        for row_idx, row in enumerate(rows):
            for col_idx, cell_val in enumerate(row):
                cell = table.rows[row_idx + 1].cells[col_idx]
                cell.text = ""
                para = cell.paragraphs[0]
                run = para.add_run(str(cell_val))
                run.font.size = Pt(8)
                run.font.color.rgb = self.colors["text"]

                # Alternate row coloring
                if row_idx % 2 == 1:
                    shading = parse_xml(
                        f'<w:shd {nsdecls("w")} w:fill="{self.config.get("light_bg", "#F3F4F6").lstrip("#")}"/>'
                    )
                    cell._tc.get_or_add_tcPr().append(shading)

        doc.add_paragraph()  # Spacer

    def generate_from_buffer(
        self,
        **kwargs: Any,
    ) -> BytesIO:
        """Generate DOCX into a BytesIO buffer."""
        buffer = BytesIO()

        # Build in-memory document
        doc = Document()

        section = doc.sections[0]
        if self.config.get("orientation") == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
            new_width = section.page_height
            section.page_width = new_width
            section.page_height = section.page_width

        section.top_margin = Pt(self.config.get("margin_top", 72))
        section.bottom_margin = Pt(self.config.get("margin_bottom", 72))
        section.left_margin = Pt(self.config.get("margin_left", 72))
        section.right_margin = Pt(self.config.get("margin_right", 72))

        title = kwargs.get("title", "Report")
        self._add_header(section, title)
        self._add_footer(section)

        # Title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(self.config.get("title_font_size", 18))
        title_run.font.color.rgb = self.colors["primary"]
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if kwargs.get("subtitle"):
            sub_para = doc.add_paragraph()
            sub_run = sub_para.add_run(kwargs["subtitle"])
            sub_run.font.size = Pt(self.config.get("subtitle_font_size", 12))
            sub_run.font.color.rgb = self.colors["text"]
            sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Date
        gen_para = doc.add_paragraph()
        gen_run = gen_para.add_run(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
        gen_run.font.size = Pt(8)
        gen_run.font.color.rgb = RGBColor(128, 128, 128)
        gen_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if kwargs.get("summary_stats"):
            self._add_stats_section(doc, kwargs["summary_stats"])

        if kwargs.get("sections"):
            for sec in kwargs["sections"]:
                self._add_text_section(doc, sec)

        if kwargs.get("tables"):
            for table_data in kwargs["tables"]:
                self._add_data_table(doc, table_data)

        doc.save(buffer)
        buffer.seek(0)
        return buffer
